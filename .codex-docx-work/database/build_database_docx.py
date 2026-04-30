from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\nikit\IdeaProjects\WW_to_Dipl")
DOC_MD = ROOT / "Documentation" / "database_documentation.md"
CHECKLIST_MD = ROOT / "Documentation" / "database_requirements_checklist.md"
SCHEMA_IMAGE = Path(r"C:\Users\nikit\Downloads\Telegram Desktop\image_2026-04-29_06-30-04.png")
OUT = ROOT / "Documentation" / "weatherwear_database_documentation.docx"

ACCENT = "1F4E79"
ACCENT_DARK = "173B5C"
HEADER_BG = "1F4E79"
LIGHT_BG = "EAF2F8"
NOTE_BG = "F4F8FB"
BORDER = "9DB5C8"
TEXT = RGBColor(32, 37, 41)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_inline_runs(paragraph, text: str) -> None:
    # Supports simple markdown inline code and bold spans.
    tokens = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(42, 79, 112)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token)


def add_body_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    add_inline_runs(paragraph, text)
    return paragraph


def add_note(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    set_paragraph_shading(paragraph, NOTE_BG)
    paragraph.paragraph_format.left_indent = Cm(0.2)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    add_inline_runs(paragraph, text)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    if level == 1:
        paragraph = doc.add_heading(text, level=1)
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        paragraph.runs[0].font.size = Pt(20)
        keep_with_next(paragraph)
        return paragraph
    if level == 2:
        paragraph = doc.add_heading(text, level=2)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.runs[0].font.color.rgb = RGBColor(37, 88, 128)
        paragraph.runs[0].font.size = Pt(14)
        keep_with_next(paragraph)
        return paragraph
    paragraph = doc.add_heading(text, level=3)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.runs[0].font.color.rgb = RGBColor(54, 95, 145)
    paragraph.runs[0].font.size = Pt(12)
    keep_with_next(paragraph)
    return paragraph


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    set_table_borders(table)

    for row_idx, source_row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx in range(max_cols):
            text = source_row[col_idx] if col_idx < len(source_row) else ""
            cell = cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, HEADER_BG)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, "F7FAFC")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline_runs(paragraph, text)
            for run in paragraph.runs:
                run.font.size = Pt(8.5 if max_cols >= 4 else 9)
                if row_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            if row_idx == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif max_cols <= 3 and col_idx in (1,):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    for raw in code_lines:
        paragraph = doc.add_paragraph()
        set_paragraph_shading(paragraph, "EEF3F7")
        paragraph.paragraph_format.left_indent = Cm(0.2)
        paragraph.paragraph_format.right_indent = Cm(0.2)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(raw if raw else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(35, 58, 79)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def should_skip_mermaid_block(lang: str) -> bool:
    return lang.strip().lower() == "mermaid"


def add_schema_figure(doc: Document) -> None:
    if not SCHEMA_IMAGE.exists():
        add_note(doc, "The schema image was not found at the expected local path, so the document uses the logical relationship summary only.")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(SCHEMA_IMAGE), width=Inches(6.7))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    cap_run = caption.add_run("Figure 1. WeatherWear logical database schema reference")
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(92, 104, 112)


def add_markdown_content(doc: Document, path: Path, skip_title=True, insert_schema_image=False) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_lines: list[str] = []
    schema_inserted = False

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_table(doc, parse_table(table_lines))
            table_lines = []

    while idx < len(lines):
        line = lines[idx]

        if line.startswith("```"):
            if not in_code:
                flush_table()
                in_code = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                if not should_skip_mermaid_block(code_lang):
                    add_code_block(doc, code_lines)
                in_code = False
                code_lang = ""
                code_lines = []
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            idx += 1
            continue
        flush_table()

        stripped = line.strip()
        if not stripped:
            idx += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            hashes, heading_text = heading_match.groups()
            if skip_title and len(hashes) == 1:
                idx += 1
                continue
            level = 1 if len(hashes) == 2 else 2 if len(hashes) == 3 else 3
            add_heading(doc, heading_text, level)
            if insert_schema_image and heading_text.startswith("4. Logical Schema") and not schema_inserted:
                add_schema_figure(doc)
                schema_inserted = True
            idx += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline_runs(paragraph, stripped[2:])
            idx += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline_runs(paragraph, numbered.group(1))
            idx += 1
            continue

        if stripped.lower().startswith("the submitted image"):
            add_note(doc, stripped)
            idx += 1
            continue

        add_body_paragraph(doc, stripped)
        idx += 1

    flush_table()


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title_run = title.add_run("WeatherWear Database Documentation")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(23, 59, 92)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle.add_run("PostgreSQL OLTP schema, integrity, roles, test data, and verification")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(88, 101, 113)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="C6D5E3", size="6")
    data = [
        ("Project", "WeatherWear Backend"),
        ("Database category", "Relational transactional database (OLTP)"),
        ("DBMS", "PostgreSQL 16"),
        ("Reviewed", "2026-04-30"),
    ]
    for row_idx, (key, value) in enumerate(data):
        row = table.rows[row_idx]
        for cell in row.cells:
            set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row.cells[0], LIGHT_BG)
        row.cells[0].paragraphs[0].add_run(key).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    overview = doc.add_paragraph()
    overview.paragraph_format.space_before = Pt(22)
    overview.paragraph_format.line_spacing = 1.12
    overview.alignment = WD_ALIGN_PARAGRAPH.CENTER
    overview.add_run(
        "This document packages the database requirements evidence for the WeatherWear backend: "
        "data dictionary, logical schema, DDL/migrations, integrity rules, access roles, "
        "test data, and verification steps."
    )
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_heading(doc, "Contents", 1)
    items = [
        "Scope and DBMS choice",
        "Requirement coverage",
        "Deployment and versioning",
        "Logical schema",
        "Data dictionary",
        "Normalization and relationships",
        "Integrity and transactions",
        "Indexes",
        "Roles and security",
        "Views, masking, and reporting",
        "Test data",
        "Restrictions compliance",
        "How to verify",
        "Appendix: database requirements checklist",
    ]
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)
    doc.add_page_break()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("WeatherWear Database Documentation - reviewed 2026-04-30")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(120, 132, 140)


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    add_markdown_content(doc, DOC_MD, skip_title=True, insert_schema_image=True)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Appendix. Database Requirements Checklist", 1)
    add_markdown_content(doc, CHECKLIST_MD, skip_title=True, insert_schema_image=False)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
